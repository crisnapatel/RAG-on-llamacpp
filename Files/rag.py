#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os, sys, argparse, json, orjson, re, glob
from pathlib import Path
from typing import List, Tuple, Dict, Iterable
import numpy as np
import requests
from tqdm import tqdm
from pypdf import PdfReader

# -------------------- Config (env overridable) --------------------
PDF_DIR        = os.environ.get("PDF_DIR", str(Path.home() / "papers"))
STORE_DIR      = os.environ.get("RAG_STORE", "./rag_store")
EMBED_URL      = os.environ.get("EMBED_URL", "http://127.0.0.1:5001")
EMBED_MODEL    = os.environ.get("EMBED_MODEL", "bge-m3")
CHAT_URL       = os.environ.get("CHAT_URL",  "http://127.0.0.1:5000")
CHAT_MODEL     = os.environ.get("CHAT_MODEL", "deepseek-r1")
CHUNK_CHARS    = int(os.environ.get("CHUNK_CHARS", "1200"))
CHUNK_OVERLAP  = int(os.environ.get("CHUNK_OVERLAP", "200"))
TOP_K          = int(os.environ.get("TOP_K", "8"))
MAX_TOKENS_OUT = int(os.environ.get("MAX_TOKENS_OUT", "700"))
TEMPERATURE    = float(os.environ.get("TEMPERATURE", "0.2"))
EMBED_MAX_ITEMS = int(os.environ.get("EMBED_MAX_ITEMS", "16"))   # cap items per HTTP call
CHARS_PER_TOKEN = float(os.environ.get("CHARS_PER_TOKEN", "3.2"))# rough T5/SPM heuristic
EMBED_TIMEOUT   = int(os.environ.get("EMBED_TIMEOUT", "600"))    # big batches can be slow


Path(STORE_DIR).mkdir(parents=True, exist_ok=True)
META_PATH  = Path(STORE_DIR) / "chunks.jsonl"     # 1 line per chunk (file,page,text)
EMB_PATH   = Path(STORE_DIR) / "embeddings.npy"   # float32, shape (N, dim)
FAISS_PATH = Path(STORE_DIR) / "index.faiss"
DIM_PATH   = Path(STORE_DIR) / "dim.json"

# -------------------- FAISS --------------------
try:
    import faiss
except Exception as e:
    print("Please: pip install 'faiss-cpu' (or faiss-gpu).", file=sys.stderr)
    sys.exit(1)

def _normalize(v: np.ndarray) -> np.ndarray:
    n = np.linalg.norm(v, axis=1, keepdims=True) + 1e-7
    return v / n

def _write_index(embs: np.ndarray):
    if embs.size == 0:
        if FAISS_PATH.exists(): FAISS_PATH.unlink()  # empty index
        return
    dim = embs.shape[1]
    index = faiss.IndexFlatIP(dim)
    index.add(_normalize(embs))
    faiss.write_index(index, str(FAISS_PATH))

def _read_index() -> faiss.IndexFlatIP | None:
    if not FAISS_PATH.exists(): return None
    return faiss.read_index(str(FAISS_PATH))

# -------------------- Extraction --------------------
def _clean_text(s: str) -> str:
    return " ".join((s or "").split())

def extract_pdf(path: Path) -> List[Dict]:
    out: List[Dict] = []
    # 1) fast path: PDFium text extraction
    try:
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(str(path))
        n = len(pdf)
        for pidx in range(n):
            page = pdf[pidx]
            txtpage = page.get_textpage()
            text = _clean_text(txtpage.get_text_range())
            if not text:
                continue
            i = 0
            while i < len(text):
                seg = text[i:i+CHUNK_CHARS].strip()
                if seg:
                    out.append({"file": str(path), "page": pidx + 1, "text": seg})
                i += max(1, CHUNK_CHARS - CHUNK_OVERLAP)
        if out:
            return out
    except Exception as e:
        print(f"[warn] pdfium parse fail {path}: {e}", file=sys.stderr)

    # 2) fallback: pypdf (more lenient, slower)
    try:
        r = PdfReader(str(path), strict=False)
        if getattr(r, "is_encrypted", False):
            try:
                ok = r.decrypt("")  # 0/1/2
                if not ok:
                    print(f"[warn] encrypted PDF (no password): {path}", file=sys.stderr)
                    return out
            except Exception as de:
                print(f"[warn] cannot decrypt {path}: {de}", file=sys.stderr)
                return out

        for pidx in range(len(r.pages)):
            try:
                page = r.pages[pidx]
                text = page.extract_text() or ""
            except Exception:
                text = ""
            text = _clean_text(text)
            if not text:
                continue
            i = 0
            while i < len(text):
                seg = text[i:i+CHUNK_CHARS].strip()
                if seg:
                    out.append({"file": str(path), "page": pidx + 1, "text": seg})
                i += max(1, CHUNK_CHARS - CHUNK_OVERLAP)
        return out
    except Exception as e:
        print(f"[warn] open/parse PDF {path}: {e}", file=sys.stderr)
        return out



# extra types
from docx import Document
import chardet

def extract_docx(path: Path) -> List[Dict]:
    try:
        doc = Document(str(path))
        text = _clean_text("\n".join(p.text for p in doc.paragraphs))
    except Exception as e:
        print(f"[warn] parse DOCX {path}: {e}", file=sys.stderr); return []
    return _chunks_from_string(text, str(path))

def extract_txt_like(path: Path) -> List[Dict]:
    try:
        raw = open(path, "rb").read()
        enc = chardet.detect(raw)["encoding"] or "utf-8"
        text = _clean_text(raw.decode(enc, errors="ignore"))
    except Exception as e:
        print(f"[warn] parse TEXT {path}: {e}", file=sys.stderr); return []
    return _chunks_from_string(text, str(path))

def _chunks_from_string(text: str, file_path: str, page_like: int = 1) -> List[Dict]:
    out, i = [], 0
    while i < len(text):
        seg = text[i:i+CHUNK_CHARS].strip()
        if seg: out.append({"file": file_path, "page": page_like, "text": seg})
        i += max(1, CHUNK_CHARS - CHUNK_OVERLAP)
    return out

def extract_file(path: Path) -> List[Dict]:
    ext = path.suffix.lower()
    if ext == ".pdf":  return extract_pdf(path)
    if ext == ".docx": return extract_docx(path)
    if ext in {".txt", ".md"}: return extract_txt_like(path)
    return []

# -------------------- Embeddings --------------------
# This does three things:
# Reads server limits (n_ctx_per_seq and n_ubatch) and packs requests to fit both.
# Automatically splits a batch and retries if the server still complains or disconnects.
# Increases timeout for big batches.

def embed_texts(texts: List[str]) -> np.ndarray:
    # empty fast-path (keep shape if dim known)
    if not texts:
        dim = int(orjson.loads(open(DIM_PATH, "rb").read())["dim"]) if DIM_PATH.exists() else 0
        return np.zeros((0, dim), dtype=np.float32) if dim else np.zeros((0, 0), dtype=np.float32)

    # discover server limits
    try:
        props = requests.get(f"{EMBED_URL}/props", timeout=5).json()
    except Exception:
        props = {}

    # ubatch tokens (encoder needs: total tokens per call <= n_ubatch)
    ubatch_limit = int(
        props.get("n_ubatch")
        or props.get("ubatch")
        or os.environ.get("EMBED_UBATCH_LIMIT", "4096")
    )

    # per-slot context tokens (total tokens per call must also fit the slot)
    # prefer explicit; else approximate ctx_size / parallel if available
    ctx_size = int(props.get("n_ctx") or os.environ.get("CTX_SIZE", "8192"))
    parallel = int(props.get("n_seq_max") or props.get("parallel") or os.environ.get("PARALLEL", "1"))
    ctx_per_seq = int(props.get("n_ctx_per_seq") or props.get("n_ctx_slot") or (ctx_size // max(1, parallel)))
    # safety margins
    token_cap = int(min(ubatch_limit, ctx_per_seq) * float(os.environ.get("PACK_SAFETY", "0.85")))

    max_items = EMBED_MAX_ITEMS
    chars_per_tok = CHARS_PER_TOKEN

    def est_tokens(s: str) -> int:
        # conservative SentencePiece-ish estimate
        return max(1, int(len(s) / chars_per_tok) + 1)

    # pack texts so sum(est_tokens) <= token_cap AND count <= max_items
    def make_batches(items: List[str]):
        batch, toks = [], 0
        for t in items:
            et = est_tokens(t)
            # if adding this would spill either limit, flush
            if batch and (toks + et > token_cap or len(batch) >= max_items):
                yield batch
                batch, toks = [], 0
            # if a single item is bigger than cap, send it alone
            if et > token_cap and not batch:
                yield [t]
                continue
            batch.append(t)
            toks += et
        if batch:
            yield batch

    results: List[np.ndarray] = []

    # recursively split a batch on server errors until it fits
    def send_batch(batch: List[str]):
        payload = {"input": batch, "model": EMBED_MODEL}
        try:
            r = requests.post(f"{EMBED_URL}/v1/embeddings", json=payload, timeout=EMBED_TIMEOUT)
            r.raise_for_status()
            vecs = [np.array(d["embedding"], dtype=np.float32) for d in r.json()["data"]]
            results.append(np.vstack(vecs))
        except (requests.ConnectionError, requests.Timeout) as e:
            if len(batch) == 1:
                raise
            mid = max(1, len(batch) // 2)
            send_batch(batch[:mid]); send_batch(batch[mid:])
        except requests.HTTPError as e:
            body = getattr(e.response, "text", "") if getattr(e, "response", None) else ""
            # split on capacity errors: 400 exceed_context_size, 413, 500, or ggml messages
            if len(batch) > 1 and (
                (getattr(e, "response", None) and e.response.status_code in (400, 413, 500))
                or "max context" in body.lower()
                or "n_ubatch" in body.lower()
                or "larger than the max" in body.lower()
            ):
                mid = max(1, len(batch) // 2)
                send_batch(batch[:mid]); send_batch(batch[mid:])
            else:
                print("[embed] server said:", body[:500], file=sys.stderr)
                raise

    batches = list(make_batches(texts))
    for b in tqdm(batches, desc="Embedding", unit="batch"):
        send_batch(b)

    embs = np.vstack(results) if results else np.zeros((0, 0), dtype=np.float32)
    if embs.size:
        with open(DIM_PATH, "wb") as f:
            f.write(orjson.dumps({"dim": int(embs.shape[1])}))
    return embs


# -------------------- Store (metas/embeddings/index) --------------------
def _load_metas() -> List[Dict]:
    if not META_PATH.exists(): return []
    return [orjson.loads(x) for x in open(META_PATH, "rb")]

def _load_embs() -> np.ndarray:
    if not EMB_PATH.exists(): return np.zeros((0, orjson.loads(open(DIM_PATH,"rb").read())["dim"]) if DIM_PATH.exists() else (0,0), dtype=np.float32)
    return np.load(EMB_PATH)

def _save_store(metas: List[Dict], embs: np.ndarray):
    # align lengths (safety)
    N = min(len(metas), embs.shape[0])
    metas = metas[:N]
    embs  = embs[:N, :] if embs.size else embs
    # write
    with open(META_PATH, "wb") as f:
        for c in metas: f.write(orjson.dumps(c) + b"\n")
    np.save(EMB_PATH, embs)
    _write_index(embs)

def _align_store() -> Tuple[int, int]:
    metas = _load_metas()
    embs  = _load_embs()
    Nm, Ne = len(metas), embs.shape[0]
    if Nm == Ne:
        # also rebuild faiss if missing or count mismatch
        idx = _read_index()
        if (idx is None) or (idx.ntotal != Ne):
            _write_index(embs)
        return 0, Ne
    # trim to min
    N = min(Nm, Ne)
    _save_store(metas[:N], embs[:N, :])
    return (max(Nm, Ne) - N, N)

def _files_in(p: str) -> List[str]:
    P = Path(p)
    if P.is_file(): return [str(P)]
    pats = ("**/*.pdf", "**/*.docx", "**/*.txt", "**/*.md")
    files = sorted(set(sum([glob.glob(os.path.join(p, pat), recursive=True) for pat in pats], [])))
    return files

def list_indexed_files() -> List[str]:
    return sorted(set(m["file"] for m in _load_metas()))

def remove_files(paths: List[str]) -> Tuple[int, int]:
    metas = _load_metas()
    embs  = _load_embs()
    if not metas: return (0, 0)
    drop = set(paths)
    keep_idx = [i for i, m in enumerate(metas) if m["file"] not in drop]
    # guard against metas/embs mismatch
    if len(metas) != embs.shape[0]:
        _align_store()
        metas = _load_metas(); embs = _load_embs()
        keep_idx = [i for i, m in enumerate(metas) if m["file"] not in drop]
    embs2 = embs[keep_idx, :] if embs.size else embs
    metas2 = [metas[i] for i in keep_idx]
    removed = len(metas) - len(metas2)
    _save_store(metas2, embs2)
    return removed, len(metas2)

def gc_missing() -> Tuple[int, int]:
    removed, remain = _align_store()
    return removed, remain

# -------------------- Indexing (append / rebuild) --------------------
from concurrent.futures import ThreadPoolExecutor, as_completed
# Parallel parse pdfs/docs (this reads PARSE_WORKERS=4 env var)
# example run command to use parallel parsing: PARSE_WORKERS=4 python rag.py index --dir "$PDF_DIR"
def _parse_files(files: Iterable[str]) -> List[Dict]:
    files = list(files)
    workers = int(os.environ.get("PARSE_WORKERS", "1"))  # default single-threaded
    if workers <= 1:
        chunks: List[Dict] = []
        for p in tqdm(files, desc="Parsing docs"):
            try:
                chunks += extract_file(Path(p))
            except Exception as e:
                print(f"[warn] {p}: {e}", file=sys.stderr)
        return chunks

    chunks: List[Dict] = []
    with ThreadPoolExecutor(max_workers=workers) as ex:
        futs = {ex.submit(extract_file, Path(p)): p for p in files}
        for fut in tqdm(as_completed(futs), total=len(futs), desc="Parsing docs"):
            p = futs[fut]
            try:
                chunks += fut.result()
            except Exception as e:
                print(f"[warn] {p}: {e}", file=sys.stderr)
    return chunks

# Single-threaded parse (backup code)
#def _parse_files(files: Iterable[str]) -> List[Dict]:
#    chunks: List[Dict] = []
#    for p in tqdm(list(files), desc="Parsing docs"):
#        try:
#            chunks += extract_file(Path(p))
#        except Exception as e:
#            print(f"[warn] {p}: {e}", file=sys.stderr)
#            # continue; do not re-raise
#    return chunks


def build_index(target_path: str, mode: str = "append") -> Tuple[int, int, int]:
    """
    mode = 'append'  -> keep existing, add new
         = 'rebuild' -> remove anything under target_path, then add fresh
    Returns: (new_chunks, nfiles_ingested, total_chunks_after)
    """
    files = _files_in(target_path)
    if not files:
        raise FileNotFoundError(f"No files under {target_path}")
    metas = _load_metas(); embs = _load_embs()
    # For rebuild: drop anything whose file path is inside target_path (or equal)
    if mode == "rebuild":
        drop = set(files) if Path(target_path).is_file() else {f for f in list_indexed_files() if str(f).startswith(str(Path(target_path)))}
        keep_idx = [i for i, m in enumerate(metas) if m["file"] not in drop]
        metas = [metas[i] for i in keep_idx]
        embs  = embs[keep_idx, :] if embs.size else embs

    chunks = _parse_files(files)
    if not chunks:
        _save_store(metas, embs)  # still rewrite index to stay consistent
        return 0, 0, len(metas)

    texts = [c["text"] for c in chunks]
    new_embs = embed_texts(texts)
    dim = int(new_embs.shape[1])
    with open(DIM_PATH, "wb") as f:
        f.write(orjson.dumps({"dim": dim}))

    metas += chunks
    embs = np.vstack([embs, new_embs]) if embs.size else new_embs
    _save_store(metas, embs)
    return len(chunks), len(files), len(metas)

# -------------------- Retrieval (dense + BM25) --------------------
from rank_bm25 import BM25Okapi
_BM25 = None
_TOKS = None

def _tokenize(s: str) -> List[str]:
    return re.findall(r"[A-Za-z0-9_µα-ωΑ-Ω+-]+", s.lower())

def _rebuild_bm25():
    global _BM25, _TOKS
    metas = _load_metas()
    _TOKS = [_tokenize(m["text"]) for m in metas]
    _BM25 = BM25Okapi(_TOKS) if _TOKS else None

def load_store():
    _align_store()
    if (_BM25 is None): _rebuild_bm25()
    idx = _read_index()
    metas = _load_metas()
    dim = orjson.loads(open(DIM_PATH,"rb").read())["dim"] if DIM_PATH.exists() else 0
    return idx, metas, dim

def retrieve(query: str, k: int) -> List[Tuple[float, Dict]]:
    idx, metas, _ = load_store()
    if (idx is None) or (idx.ntotal == 0) or (not metas):
        return []
    # dense
    q = embed_texts([query])
    D, I = idx.search(_normalize(q), max(k*3, 24))
    dense_hits = [(float(D[0][j]), int(I[0][j])) for j in range(len(I[0])) if I[0][j] >= 0]
    # bm25
    bm = _BM25.get_scores(_tokenize(query)) if _BM25 else np.zeros(len(metas))
    bm_z = (bm - bm.mean()) / (bm.std() + 1e-6) if len(metas) else bm
    scored = {}
    for s, idxi in dense_hits:
        scored[idxi] = scored.get(idxi, 0.0) + s
    for idxi, b in enumerate(bm_z):
        if idxi in scored:
            scored[idxi] = scored[idxi] + 0.2 * float(b)
    top = sorted(scored.items(), key=lambda x: x[1], reverse=True)[:k]
    return [(score, metas[idxi]) for idxi, score in top]

# -------------------- Chat --------------------
def _strip_think(s: str) -> str:
    return re.sub(r"<think>.*?</think>\s*", "", s, flags=re.DOTALL).strip()

def _escape_html(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

def chat_complete(system: str, user: str, raw: bool) -> str:
    payload = {
        "model": CHAT_MODEL,
        "messages": [{"role":"system","content":system},
                     {"role":"user","content":user}],
        "temperature": TEMPERATURE,
        "max_tokens": MAX_TOKENS_OUT,
        "stream": False,
        "add_generation_prompt": True
    }
    r = requests.post(f"{CHAT_URL}/v1/chat/completions", json=payload, timeout=180)
    r.raise_for_status()
    msg = (r.json()["choices"][0]["message"]["content"] or "").strip()
    if raw:
        # show <think> safely
        m = re.search(r"<think>(.*?)</think>", msg, flags=re.DOTALL)
        think = m.group(1).strip() if m else ""
        final = _strip_think(msg)
        if think:
            return (
                "### Raw model output\n"
                "<details><summary>Reasoning &lt;think&gt; (click to expand)</summary>\n\n"
                f"```\n{think}\n```\n"
                "</details>\n\n"
                f"{final}"
            ).strip()
        # if model output is only <think> (rare), still show it
        return "```\n" + _escape_html(msg) + "\n```"
    # normal: strip think
    clean = _strip_think(msg)
    return clean or "The model returned an empty message."

SYS_PROMPT = (
  "You are a precise scientific research assistant."
  "Answer ONLY from the provided context. If unknown, say you don't know."
  "Use short inline citations like [#] matching the Sources list. If there are no citations, do not cite."
)

def make_prompt(q: str, hits: List[Tuple[float,Dict]]):
    ctx, sources = [], []
    for j,(score,m) in enumerate(hits,1):
        h = f"[{j}] {Path(m['file']).name}, p.{m['page']}"
        ctx.append(f"{h}\n{m['text']}\n")
        sources.append(h)
    blob = "\n\n---\n\n".join(ctx) if ctx else "(no context)"
    user = (
        f"Question:\n{q}\n\nContext:\n{blob}\n\n"
        "Instructions:\n- Use ONLY the context.\n- Cite with [#].\n- Be concise."
    )
    return user, sources

# -------------------- CLI --------------------
def cmd_index(a):
    chunks, nfiles, total = build_index(a.dir, mode="append")
    print(f"Appended {chunks} chunks from {nfiles} files. Total: {total}")

def cmd_ask(a):
    hits = retrieve(a.question, a.k)
    user, src = make_prompt(a.question, hits)
    ans = chat_complete(SYS_PROMPT, user, raw=False)
    print("\n=== Answer ===\n" + ans.strip() + "\n\n=== Sources ===")
    for s in src: print(s)

# -------------------- UI --------------------
def cmd_ui(a):
    import gradio as gr

    def _stats():
        _align_store()
        metas = _load_metas()
        files = sorted(set(m["file"] for m in metas))
        return f"{len(metas)} chunks • {len(files)} files"

    def _dropdown_choices():
        return list_indexed_files()

    def do_append_folder(folder):
        chunks, nfiles, total = build_index(folder, mode="append")
        _rebuild_bm25()
        return (f"Appended ✔ {chunks} chunks from {nfiles} files. Total: {total}",
                _stats(),
                gr.update(choices=_dropdown_choices(), value=[]))

    def do_rebuild_folder(folder):
        chunks, nfiles, total = build_index(folder, mode="rebuild")
        _rebuild_bm25()
        return (f"Rebuilt ✔ {chunks} chunks from {nfiles} files. Total: {total}",
                _stats(),
                gr.update(choices=_dropdown_choices(), value=[]))

    def do_append_file(path):
        chunks, nfiles, total = build_index(path, mode="append")
        _rebuild_bm25()
        return (f"Appended ✔ {chunks} chunks from {nfiles} file. Total: {total}",
                _stats(),
                gr.update(choices=_dropdown_choices(), value=[]))

    def do_rebuild_file(path):
        chunks, nfiles, total = build_index(path, mode="rebuild")
        _rebuild_bm25()
        return (f"Rebuilt ✔ {chunks} chunks from {nfiles} file. Total: {total}",
                _stats(),
                gr.update(choices=_dropdown_choices(), value=[]))

    def do_remove(paths):
        try:
            removed, remain = remove_files(paths or [])
            _rebuild_bm25()
            return (f"Removed ✔ {removed} chunks. Remaining: {remain}",
                    _stats(),
                    gr.update(choices=_dropdown_choices(), value=[]))
        except Exception as e:
            return (f"Error: {e}", _stats(), gr.update(choices=_dropdown_choices(), value=[]))

    def do_gc():
        try:
            removed, remain = gc_missing()
            _rebuild_bm25()
            return (f"GC ✔ removed {removed} chunks. Remaining: {remain}",
                    _stats(),
                    gr.update(choices=_dropdown_choices(), value=[]))
        except Exception as e:
            return (f"Error: {e}", _stats(), gr.update(choices=_dropdown_choices(), value=[]))

    def qa(msg, history, k, raw):
        q = msg if isinstance(msg, str) else (msg.get("content","") if isinstance(msg, dict) else str(msg))
        hits = retrieve(q, int(k))
        user, src = make_prompt(q, hits)
        ans = chat_complete(SYS_PROMPT, user, raw=bool(raw))
        cite = "\n\n---\n" + "\n".join(src) if src else ""
        return ans + cite

    with gr.Blocks(title="Local PDF RAG (llama.cpp)", css="""
    .wrap {max-width: 1280px; margin: 0 auto;}
    .bigbot {height:520px;}
    """) as app:
        with gr.Row(elem_classes="wrap"):
            # Left: corpus manager
            with gr.Column(scale=4, min_width=320):
                gr.Markdown("### Corpus manager")
                folder_tb = gr.Textbox(value=os.environ.get("PDF_DIR", PDF_DIR), label="Folder")
                with gr.Row():
                    btn_append_folder  = gr.Button("Append folder")
                    btn_rebuild_folder = gr.Button("Rebuild from folder")
                file_tb = gr.Textbox(value="", label="Single file (PDF/DOCX/TXT/MD)")
                with gr.Row():
                    btn_append_file = gr.Button("Append file")
                    btn_rebuild_file = gr.Button("Rebuild from file")

                status = gr.Markdown(_stats())

                gr.Markdown("### Indexed files (select to remove)")
                files_dd = gr.Dropdown(choices=_dropdown_choices(), multiselect=True, value=[], label="")
                with gr.Row():
                    btn_remove = gr.Button("Remove selected")
                    btn_gc     = gr.Button("GC missing")

                sidebar_msg = gr.Markdown("")

            # Right: search & chat
            with gr.Column(scale=8, min_width=640):
                gr.Markdown("### Search settings")
                with gr.Row():
                    k = gr.Slider(3, 24, value=float(TOP_K), step=1, label="Top-K passages (retrieval)")
                    show_raw = gr.Checkbox(value=False, label="Show raw model output (may include <think>)")
                chat = gr.ChatInterface(
                    fn=lambda message, history: qa(message, history, k.value, show_raw.value),
                    type="messages",
                )

        # wire buttons
        btn_append_folder.click(fn=do_append_folder, inputs=folder_tb, outputs=[sidebar_msg, status, files_dd])
        btn_rebuild_folder.click(fn=do_rebuild_folder, inputs=folder_tb, outputs=[sidebar_msg, status, files_dd])
        btn_append_file.click(fn=do_append_file, inputs=file_tb, outputs=[sidebar_msg, status, files_dd])
        btn_rebuild_file.click(fn=do_rebuild_file, inputs=file_tb, outputs=[sidebar_msg, status, files_dd])
        btn_remove.click(fn=do_remove, inputs=files_dd, outputs=[sidebar_msg, status, files_dd])
        btn_gc.click(fn=do_gc, inputs=None, outputs=[sidebar_msg, status, files_dd])

        app.queue().launch(server_name="0.0.0.0", server_port=7860)

# -------------------- argparse --------------------
if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    sub = ap.add_subparsers()

    p0 = sub.add_parser("index")
    p0.add_argument("--dir", default=PDF_DIR)
    p0.set_defaults(func=cmd_index)

    p1 = sub.add_parser("ask")
    p1.add_argument("question")
    p1.add_argument("-k","--k",type=int,default=TOP_K)
    p1.set_defaults(func=cmd_ask)

    p2 = sub.add_parser("ui")
    p2.set_defaults(func=cmd_ui)

    a = ap.parse_args()
    a.func(a) if hasattr(a,"func") else ap.print_help()

